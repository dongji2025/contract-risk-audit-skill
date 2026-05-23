#!/usr/bin/env python3
"""
Enhanced Word Report Generator — Comprehensive Contract Risk Audit Report
基于《民法典》第463-647条 + 《商业秘密保护规定》的图文并茂Word报告

Features:
- Professional title page with risk level badge
- Executive summary with risk statistics
- Multiple chart visualizations (pie, bar, radar, gauge)
- Company information section
- Risk overview with category breakdown matrix
- Detailed clause-by-clause analysis with full legal references
- Trade secret / confidentiality special analysis
- High risk items with detailed modification suggestions
- Civil Code article cross-reference table
- Prioritized recommendations
- Full legal reference index
- Professional disclaimer
"""

import sys
import json
import argparse
from datetime import datetime
from io import BytesIO
import os

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm

# Late import for docx (module-level needed for helper functions)
try:
    from docx.shared import Inches, Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    pass


# ── CJK Font Detection ───────────────────────────
def _detect_cjk_font():
    preferred = [
        'Microsoft YaHei', 'SimHei', 'PingFang SC',
        'Noto Sans CJK SC', 'WenQuanYi Micro Hei',
        'WenQuanYi Zen Hei', 'Arial Unicode MS', 'DejaVu Sans', 'sans-serif'
    ]
    available = {f.name for f in fm.fontManager.ttflist}
    for font in preferred:
        if font in available:
            return [font] + [f for f in preferred if f != font]
    return preferred


plt.rcParams['font.sans-serif'] = _detect_cjk_font()
plt.rcParams['axes.unicode_minus'] = False

# ── Color palette ─────────────────────────────────
COLOR_HIGH = '#DC143C'
COLOR_MEDIUM = '#FFA500'
COLOR_LOW = '#228B22'
COLOR_CRITICAL = '#8B0000'
COLOR_PRIMARY = '#003366'
COLOR_ACCENT = '#667eea'
COLOR_GREY = '#808080'
COLOR_LIGHT_BG = '#f8f9fa'


# ══════════════════════════════════════════════════
# Chart generation functions
# ══════════════════════════════════════════════════
def create_pie_chart(data, labels, title, colors=None):
    if colors is None:
        colors = [COLOR_HIGH, COLOR_MEDIUM, COLOR_LOW, COLOR_ACCENT]
    fig, ax = plt.subplots(figsize=(5.5, 5.5))
    wedges, texts, autotexts = ax.pie(
        data, labels=labels, autopct='%1.1f%%',
        colors=colors, startangle=90,
        textprops={'fontsize': 11}
    )
    for at in autotexts:
        at.set_color('white')
        at.set_fontweight('bold')
        at.set_fontsize(10)
    ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()
    return buf


def create_stacked_bar(categories, high_data, medium_data, low_data, title):
    fig, ax = plt.subplots(figsize=(10, 5))
    x = range(len(categories))
    w = 0.6
    ax.bar(x, high_data, w, label='高风险', color=COLOR_HIGH)
    ax.bar(x, medium_data, w, bottom=high_data, label='中风险', color=COLOR_MEDIUM)
    ax.bar(x, low_data, w, bottom=[h + m for h, m in zip(high_data, medium_data)], label='低风险', color=COLOR_LOW)
    ax.set_ylabel('条款数量', fontsize=11)
    ax.set_title(title, fontsize=14, fontweight='bold')
    ax.set_xticks(x)
    ax.set_xticklabels(categories, rotation=45, ha='right', fontsize=8)
    ax.legend(loc='upper right', fontsize=9)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()
    return buf


def create_hbar_chart(category_data, title):
    categories = list(category_data.keys())
    values = list(category_data.values())
    fig, ax = plt.subplots(figsize=(7, 5))
    colors = [COLOR_HIGH if v >= 25 else COLOR_MEDIUM if v >= 15 else COLOR_LOW for v in values]
    bars = ax.barh(categories, values, color=colors, edgecolor='white', height=0.65)
    for bar, val in zip(bars, values):
        ax.text(val + 0.5, bar.get_y() + bar.get_height() / 2, f'{val}', va='center', fontsize=9, fontweight='bold')
    ax.set_xlabel('风险评分', fontsize=11)
    ax.set_title(title, fontsize=14, fontweight='bold')
    ax.set_xlim(0, max(values) + 15)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()
    return buf


def create_risk_gauge(score, title):
    fig, ax = plt.subplots(figsize=(5, 3.5), subplot_kw={'projection': 'polar'})
    ax.bar(0, 0.25, width=0.5, bottom=0.75, color=COLOR_LOW, alpha=0.25)
    ax.bar(0.5, 0.25, width=0.5, bottom=0.75, color=COLOR_MEDIUM, alpha=0.25)
    ax.bar(1.0, 0.25, width=0.5, bottom=0.75, color=COLOR_HIGH, alpha=0.25)
    ax.set_ylim(0, 1)
    ax.set_yticklabels([])
    ax.set_xticklabels([])
    ax.spines['polar'].set_visible(False)
    ax.grid(False)
    ax.annotate('低', xy=(0.25, 0.88), fontsize=9, ha='center', color=COLOR_LOW)
    ax.annotate('中', xy=(0.75, 0.88), fontsize=9, ha='center', color=COLOR_MEDIUM)
    ax.annotate('高', xy=(1.25, 0.88), fontsize=9, ha='center', color=COLOR_HIGH)
    pointer_theta = min((score / 100) * 1.5 - 0.25, 1.25)
    ax.arrow(pointer_theta, 0.3, 0, 0.5, head_width=0.1, head_length=0.05,
             fc='black', ec='black', alpha=0.85)
    ax.set_title(f'{title}\n风险评分: {score} 分', fontsize=11, fontweight='bold', pad=8)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()
    return buf


def create_radar_chart(categories, values, title):
    N = len(categories)
    if N < 3:
        return None
    angles = [n / float(N) * 2 * 3.1415926535 for n in range(N)]
    angles += angles[:1]
    values += values[:1]

    fig, ax = plt.subplots(figsize=(6, 6), subplot_kw={'projection': 'polar'})
    ax.fill(angles, values, alpha=0.25, color=COLOR_ACCENT)
    ax.plot(angles, values, 'o-', linewidth=2, color=COLOR_ACCENT)
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=8)
    ax.set_ylim(0, max(values) * 1.2)
    ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
    plt.tight_layout()
    buf = BytesIO()
    plt.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close()
    return buf


# ══════════════════════════════════════════════════
# Main report generator
# ══════════════════════════════════════════════════
def generate_report(risk_data_path, output_path=None):
    """Generate comprehensive Word report with charts, legal references, and recommendations."""
    try:
        from docx import Document
        from docx.shared import Inches, Cm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        return 1

    # ── Load data ──
    try:
        with open(risk_data_path, 'r', encoding='utf-8') as f:
            risk_data = json.load(f)
    except Exception as e:
        print(f"ERROR reading risk data: {e}")
        return 1

    contract_info = risk_data.get('contract_info', {})
    summary = risk_data.get('summary', {})
    clauses = risk_data.get('clauses', [])
    high_risk = risk_data.get('high_risk_clauses', [])
    medium_risk = risk_data.get('medium_risk_clauses', [])
    low_risk = risk_data.get('low_risk_clauses', [])
    trade_secret = risk_data.get('trade_secret_analysis', {})
    # Fix: findings are inside trade_secret_analysis, not top-level
    trade_findings = trade_secret.get('findings', []) if trade_secret else []
    all_refs = risk_data.get('all_legal_references', [])

    # ── Create document ──
    doc = Document()
    _setup_page(doc)

    # ── Build sections ──
    add_cover_page(doc, summary, contract_info)
    doc.add_page_break()
    add_table_of_contents(doc)
    doc.add_page_break()
    add_executive_summary(doc, summary, risk_data)
    add_chart_section(doc, summary, clauses)
    add_company_info_section(doc, contract_info, risk_data)
    add_risk_overview_table(doc, summary, clauses)
    add_trade_secret_section(doc, trade_secret, trade_findings)
    add_clause_detail_section(doc, risk_data)
    add_civil_code_reference_table(doc, clauses)
    add_priority_recommendations(doc, risk_data)
    add_legal_reference_index(doc, all_refs)
    add_disclaimer_final(doc)
    add_footer(doc)

    # ── Save ──
    if not output_path:
        output_path = "contract_risk_audit_report.docx"
    doc.save(output_path)
    print(f"[OK] Word report generated: {output_path}")
    return 0


# ══════════════════════════════════════════════════
# Page setup
# ══════════════════════════════════════════════════
def _setup_page(doc):
    from docx.shared import Inches
    for section in doc.sections:
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)


# ══════════════════════════════════════════════════
# Table of Contents
# ══════════════════════════════════════════════════
def add_table_of_contents(doc):
    """Add a visual table of contents page"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('目  录')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('TABLE OF CONTENTS')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 120, 120)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()
    doc.add_paragraph()

    toc_items = [
        ('一、', '执行摘要', '合同风险审查总体结论与关键指标'),
        ('二、', '风险可视化分析', '饼图、仪表盘、柱状图、排名图'),
        ('三、', '合同主体信息', '甲乙方公司信息及资质核验'),
        ('四、', '风险概览', '风险类别分布矩阵与统计'),
        ('五、', '商业秘密与保密条款专项审查', '保密范围、期限、竞业限制、保密措施'),
        ('六、', '条款详细分析', '逐条风险识别、法律依据、修改建议'),
        ('七、', '《民法典》条款覆盖清单', '审查涉及的法条索引'),
        ('八、', '优先修订建议', '必须修订/建议协商/通用建议'),
        ('九、', '法律依据索引', '全部引用法律法规汇总'),
        ('十、', '免责声明', '法律效力说明'),
    ]

    for num, title, desc in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'  {num}{title}')
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0, 51, 102)
        run = p.add_run(f'\n      {desc}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(140, 140, 140)
        doc.add_paragraph()

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)


# ══════════════════════════════════════════════════
# Cover Page
# ══════════════════════════════════════════════════
def add_cover_page(doc, summary, contract_info):
    """Professional cover page with rich visual elements"""

    # ── Top decorative bar ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)

    # ── Report classification ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('【法律文件·机密】')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(180, 0, 0)
    run.bold = True

    doc.add_paragraph()  # spacer

    # ── Main title block with decorative rules ──
    # Upper rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('')

    # Main Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('合 同 风 险 审 查 报 告')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('CONTRACT RISK AUDIT REPORT')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(120, 120, 120)

    # Lower rule
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('')
    run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    # ── Legal basis badge ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('┌─────────────────────────────────────────┐')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('│  审查依据：《中华人民共和国民法典》第463-647条  │')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('│  《商业秘密保护规定》 · 条款级逐条比对分析  │')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('└─────────────────────────────────────────┘')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Risk Level Badge (boxed) ──
    risk_level = summary.get('risk_level', 'UNKNOWN')
    level_colors = {
        'CRITICAL': RGBColor(139, 0, 0),
        'HIGH': RGBColor(200, 50, 50),
        'MEDIUM': RGBColor(200, 150, 0),
        'LOW': RGBColor(0, 128, 0),
    }
    level_labels = {
        'CRITICAL': '严重风险',
        'HIGH': '高风险',
        'MEDIUM': '中等风险',
        'LOW': '低风险',
    }
    level_icons = {
        'CRITICAL': '🔴',
        'HIGH': '🟠',
        'MEDIUM': '🟡',
        'LOW': '🟢',
    }
    lc = level_colors.get(risk_level, RGBColor(100, 100, 100))
    ll = level_labels.get(risk_level, risk_level)
    li = level_icons.get(risk_level, '⚪')

    # Risk badge box top
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('┌──────────────────────────────┐')
    run.font.size = Pt(8)
    run.font.color.rgb = lc
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'│   综合风险等级              │')
    run.font.size = Pt(9)
    run.font.color.rgb = lc
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'│   {li} {ll} ({risk_level})   │')
    run.font.size = Pt(22)
    run.font.color.rgb = lc
    run.bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'│   综合风险评分：{summary.get("total_risk_score", 0)} 分        │')
    run.font.size = Pt(10)
    run.font.color.rgb = lc
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('└──────────────────────────────┘')
    run.font.size = Pt(8)
    run.font.color.rgb = lc

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Key stats box ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('┌──────────────────────────────────────────────┐')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'│  总条款：{summary.get("total_clauses", 0)} 条  │  '
                     f'高风险：{summary.get("high_risk_count", 0)} 条  │  '
                     f'中风险：{summary.get("medium_risk_count", 0)} 条  │  '
                     f'低风险：{summary.get("low_risk_count", 0)} 条  │')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'│  《民法典》条文覆盖：{summary.get("civil_code_articles_covered", 0)} 条                      │')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('└──────────────────────────────────────────────┘')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ── Bottom section ──
    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('━' * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0, 51, 102)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'合同主体：{contract_info.get("name", "未明确")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(60, 60, 60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'报告生成日期：{datetime.now().strftime("%Y年%m月%d日")}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('本报告为自动化风险分析工具生成的参考性文件，不构成法律意见')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(150, 150, 150)


# ══════════════════════════════════════════════════
# Executive Summary
# ══════════════════════════════════════════════════
def add_executive_summary(doc, summary, risk_data):
    h = doc.add_heading('一、执行摘要', 1)
    _style_heading(h)

    p = doc.add_paragraph()
    run = p.add_run(
        '本报告依据《中华人民共和国民法典》第463-647条及《商业秘密保护规定》'
        '对合同进行了全条款、逐条目的风险审查分析。分析覆盖合同效力、订立程序、'
        '履行义务、违约责任、保密条款、竞业限制、债权转让、合同解除等多个维度，'
        '每条风险均标注对应的法律依据和具体的修改建议。'
    )
    run.font.size = Pt(11)

    doc.add_paragraph()

    # Key stats table
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    data = [
        ('分析指标', '数值'),
        ('总条款数', str(summary.get('total_clauses', 0))),
        ('高风险条款', f"{summary.get('high_risk_count', 0)} 条"),
        ('中风险条款', f"{summary.get('medium_risk_count', 0)} 条"),
        ('低风险条款', f"{summary.get('low_risk_count', 0)} 条"),
        ('覆盖《民法典》条文', f"{summary.get('civil_code_articles_covered', 0)} 条"),
        ('综合风险评分', f"{summary.get('total_risk_score', 0)} 分"),
        ('综合风险等级', summary.get('risk_level', 'N/A')),
    ]

    for i, (label, value) in enumerate(data):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = label
        c1.text = value
        if i == 0:
            for c in [c0, c1]:
                for p in c.paragraphs:
                    p.runs[0].bold = True if p.runs else False

    doc.add_paragraph()

    # Warning
    hc = summary.get('high_risk_count', 0)
    if hc > 0:
        p = doc.add_paragraph()
        run = p.add_run(f'⚠ 重要提示：该合同存在 {hc} 条高风险条款，建议在签署前进行必要修订。')
        run.font.color.rgb = RGBColor(180, 0, 0)
        run.bold = True
        run.font.size = Pt(11)

    if hc >= 5:
        p = doc.add_paragraph()
        run = p.add_run('合同存在较多高风险条款，强烈建议在专业律师协助下进行全面修订后再签署。')
        run.font.color.rgb = RGBColor(139, 0, 0)
        run.bold = True
        run.font.size = Pt(11)


# ══════════════════════════════════════════════════
# Chart Section
# ══════════════════════════════════════════════════
def add_chart_section(doc, summary, clauses):
    h = doc.add_heading('二、风险可视化分析', 1)
    _style_heading(h)

    hc = summary.get('high_risk_count', 0)
    mc = summary.get('medium_risk_count', 0)
    lc = summary.get('low_risk_count', 0)

    # --- Pie Chart ---
    if hc + mc + lc > 0:
        pie_buf = create_pie_chart(
            [hc, mc, lc],
            [f'高风险({hc})', f'中风险({mc})', f'低风险({lc})'],
            '合同风险等级分布',
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(pie_buf, width=Inches(4.2))

        doc.add_paragraph()

    # --- Gauge Chart ---
    score = summary.get('total_risk_score', 0)
    gauge_buf = create_risk_gauge(score, '综合风险评分仪表盘')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(gauge_buf, width=Inches(3.8))

    doc.add_paragraph()

    # --- Category stacked bar ---
    categories = {}
    for c in clauses:
        cat = c.get('category', '其他')
        lv = c.get('risk_level', 'LOW')
        if cat not in categories:
            categories[cat] = {'HIGH': 0, 'MEDIUM': 0, 'LOW': 0}
        if lv in categories[cat]:
            categories[cat][lv] += 1

    if categories:
        cat_names = list(categories.keys())
        bar_buf = create_stacked_bar(
            cat_names,
            [categories[c]['HIGH'] for c in cat_names],
            [categories[c]['MEDIUM'] for c in cat_names],
            [categories[c]['LOW'] for c in cat_names],
            '各风险类别条款分布',
        )
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(bar_buf, width=Inches(6))

        doc.add_paragraph()

    # --- Category score ranking ---
    cat_scores = {}
    for c in clauses:
        cat = c.get('category', '其他')
        s = c.get('risk_score', 0)
        cat_scores[cat] = cat_scores.get(cat, 0) + s

    if cat_scores:
        sorted_scores = dict(sorted(cat_scores.items(), key=lambda x: x[1], reverse=True)[:10])
        hbar_buf = create_hbar_chart(sorted_scores, '风险类别评分排名 (Top 10)')
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(hbar_buf, width=Inches(5.5))

    doc.add_paragraph()


# ══════════════════════════════════════════════════
# Company Info
# ══════════════════════════════════════════════════
def add_company_info_section(doc, contract_info, risk_data=None):
    h = doc.add_heading('三、合同主体信息', 1)
    _style_heading(h)

    if not any(contract_info.values()):
        p = doc.add_paragraph()
        p.add_run('未能提取到完整的主体信息，建议补全。')
        return

    p = doc.add_paragraph()
    run = p.add_run('以下信息从合同文本中自动提取，请在签署前核实准确性。')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(120, 120, 120)
    doc.add_paragraph()

    # ── Main company info ──
    p = doc.add_paragraph()
    run = p.add_run('▎ 合同主体信息（乙方/服务提供方）')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    table = doc.add_table(rows=10, cols=2)
    table.style = 'Table Grid'

    fields = [
        ('公司名称', contract_info.get('name', '未明确')),
        ('统一社会信用代码', contract_info.get('unified_credit_code', '未明确')),
        ('法定代表人', contract_info.get('legal_representative', '未明确')),
        ('注册资本', contract_info.get('registered_capital', '未明确')),
        ('实缴资本', contract_info.get('paid_capital', '未明确')),
        ('注册地址', contract_info.get('address', '未明确')),
        ('开户银行', contract_info.get('bank_name', '未明确')),
        ('银行账号', contract_info.get('bank_account', '未明确')),
        ('联系电话', contract_info.get('phone', '未明确')),
    ]

    for i, (label, value) in enumerate(fields):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c0.text = label
        c1.text = str(value)
        for p in c0.paragraphs:
            if p.runs:
                p.runs[0].bold = True
        # Color code completeness
        if value in ('未明确', '', None):
            for p in c1.paragraphs:
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(200, 50, 50)

    doc.add_paragraph()

    # ── Verification checklist ──
    p = doc.add_paragraph()
    run = p.add_run('▎ 主体信息核验清单')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    checklist = [
        ('统一社会信用代码', '通过国家企业信用信息公示系统核实企业的工商登记状态'),
        ('法定代表人', '核实法定代表人身份是否与工商登记一致'),
        ('经营范围', '确认合同涉及的业务在对方经营范围内（涉及特许经营的须查验许可证）'),
        ('银行账户', '核实收款账户名称与合同主体名称一致，谨防诈骗'),
        ('授权代表', '签约人须持有合法有效的授权委托书'),
    ]
    for label, desc in checklist:
        p = doc.add_paragraph()
        run = p.add_run(f'  ☐ {label}：')
        run.bold = True
        run.font.size = Pt(10)
        run = p.add_run(desc)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)


# ══════════════════════════════════════════════════
# Risk Overview Table
# ══════════════════════════════════════════════════
def add_risk_overview_table(doc, summary, clauses):
    h = doc.add_heading('四、风险概览', 1)
    _style_heading(h)

    p = doc.add_paragraph()
    run = p.add_run(
        f'合同共 {summary.get("total_clauses", 0)} 条条款，'
        f'其中高风险 {summary.get("high_risk_count", 0)} 条、'
        f'中风险 {summary.get("medium_risk_count", 0)} 条、'
        f'低风险 {summary.get("low_risk_count", 0)} 条。'
        f'分析覆盖《民法典》第463-647条共 {summary.get("civil_code_articles_covered", 0)} 条条文。'
    )
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Category breakdown
    cat_data = {}
    for c in clauses:
        cat = c.get('category', '其他')
        lv = c.get('risk_level', 'LOW')
        if cat not in cat_data:
            cat_data[cat] = {'HIGH': 0, 'MEDIUM': 0, 'LOW': 0, 'total': 0, 'score': 0}
        cat_data[cat][lv] += 1
        cat_data[cat]['total'] += 1
        cat_data[cat]['score'] += c.get('risk_score', 0)

    if cat_data:
        table = doc.add_table(rows=len(cat_data) + 1, cols=6)
        table.style = 'Table Grid'
        headers = ['风险类别', '高风险', '中风险', '低风险', '小计', '风险评分']
        for i, hdr in enumerate(headers):
            c = table.rows[0].cells[i]
            c.text = hdr
            for p in c.paragraphs:
                if p.runs:
                    p.runs[0].bold = True

        for i, (cat, cnt) in enumerate(sorted(cat_data.items()), 1):
            table.rows[i].cells[0].text = cat
            table.rows[i].cells[1].text = str(cnt['HIGH'])
            table.rows[i].cells[2].text = str(cnt['MEDIUM'])
            table.rows[i].cells[3].text = str(cnt['LOW'])
            table.rows[i].cells[4].text = str(cnt['total'])
            table.rows[i].cells[5].text = str(cnt['score'])


# ══════════════════════════════════════════════════
# Trade Secret / Confidentiality Section
# ══════════════════════════════════════════════════
def add_trade_secret_section(doc, trade_secret, trade_findings):
    h = doc.add_heading('五、商业秘密与保密条款专项审查', 1)
    _style_heading(h)

    if not trade_secret:
        p = doc.add_paragraph()
        p.add_run('未提取到保密条款相关数据。')
        return

    p = doc.add_paragraph()
    run = p.add_run('本项审查依据《商业秘密保护规定》及《反不正当竞争法》第9条、第17条，'
                     '对合同中的保密条款、竞业限制、保密措施等进行全方位审查。')
    run.font.size = Pt(10)
    doc.add_paragraph()

    # ── Status overview card-style table ──
    p = doc.add_paragraph()
    run = p.add_run('▎ 保密条款合规性检查清单')
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)
    doc.add_paragraph()

    table = doc.add_table(rows=10, cols=3)
    table.style = 'Table Grid'
    # Header
    for i, hdr in enumerate(['审查项目', '状态', '评估']):
        c = table.rows[0].cells[i]
        c.text = hdr
        for p in c.paragraphs:
            if p.runs:
                p.runs[0].bold = True

    checklist = [
        ('保密条款存在', '✅ 已约定' if trade_secret.get('has_confidentiality_clause') else '❌ 缺失',
         '合规' if trade_secret.get('has_confidentiality_clause') else '须补充'),
        ('保密范围界定', '✅ 已明确' if trade_secret.get('confidentiality_scope_defined') else '❌ 未明确界定',
         '合规' if trade_secret.get('confidentiality_scope_defined') else '须补充'),
        ('保密期限约定', '✅ 已约定' if trade_secret.get('confidentiality_period_defined') else '❌ 未约定',
         '合规' if trade_secret.get('confidentiality_period_defined') else '须补充'),
        ('保密期限合理性', '✅ 合理' if trade_secret.get('confidentiality_period_reasonable') else '❌ 永久/无限期',
         '合规' if trade_secret.get('confidentiality_period_reasonable') else '不合规'),
        ('保密措施约定', '✅ 已约定措施' if trade_secret.get('confidentiality_measures_defined') else '❌ 措施不具体',
         '合规' if trade_secret.get('confidentiality_measures_defined') else '须补充'),
        ('竞业限制条款', f'有（{trade_secret.get("non_compete_period_months", 0)}个月）' if trade_secret.get('non_compete_present') else '无',
         '合规' if not trade_secret.get('non_compete_present') or trade_secret.get('non_compete_period_months', 0) <= 24 else '不合规'),
        ('竞业限制补偿', '✅ 已约定补偿' if trade_secret.get('non_compete_compensation') else '❌ 未约定补偿',
         '合规' if trade_secret.get('non_compete_compensation') or not trade_secret.get('non_compete_present') else '不合规'),
        ('返还/销毁义务', '✅ 已约定' if trade_secret.get('return_or_destroy_clause') else '❌ 缺失',
         '合规' if trade_secret.get('return_or_destroy_clause') else '须补充'),
        ('保密违约救济', '✅ 已约定' if trade_secret.get('breach_remedies_defined') else '❌ 未约定',
         '合规' if trade_secret.get('breach_remedies_defined') else '须补充'),
    ]

    for i, (label, value, assessment) in enumerate(checklist, 1):
        c0 = table.rows[i].cells[0]
        c1 = table.rows[i].cells[1]
        c2 = table.rows[i].cells[2]
        c0.text = label
        c1.text = value
        c2.text = assessment
        for p in c0.paragraphs:
            if p.runs:
                p.runs[0].bold = True
        # Color the assessment
        for p in c2.paragraphs:
            if p.runs:
                if '不合规' in assessment:
                    p.runs[0].font.color.rgb = RGBColor(180, 0, 0)
                elif '须补充' in assessment:
                    p.runs[0].font.color.rgb = RGBColor(200, 150, 0)
                else:
                    p.runs[0].font.color.rgb = RGBColor(0, 128, 0)

    doc.add_paragraph()

    # Findings
    if trade_findings:
        p = doc.add_paragraph()
        run = p.add_run('保密条款审查发现：')
        run.bold = True
        run.font.size = Pt(11)

        for finding in trade_findings:
            p = doc.add_paragraph()
            run = p.add_run(f'  ⚠ {finding}')
            run.font.size = Pt(10)
            if '永久' in finding or '超期' in finding or '无补偿' in finding:
                run.font.color.rgb = RGBColor(180, 0, 0)

    # Legal basis for trade secret review
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('商业秘密保护审查法律依据：')
    run.bold = True
    run.font.size = Pt(10)
    refs = [
        '《反不正当竞争法》第9条（商业秘密定义与侵权行为）',
        '《反不正当竞争法》第17条（赔偿责任）',
        '最高人民法院《关于审理侵犯商业秘密民事案件适用法律若干问题的规定》（法释〔2020〕7号）',
        '《民法典》第501条（缔约阶段的保密义务）',
        '《民法典》第558条（合同终止后的保密义务）',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        run = p.add_run(f'  • {ref}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(80, 80, 80)


# ══════════════════════════════════════════════════
# Detailed Clause Analysis
# ══════════════════════════════════════════════════
def add_clause_detail_section(doc, risk_data):
    """Detailed clause-by-clause analysis — all risk levels with full analysis"""
    h = doc.add_heading('六、条款详细分析', 1)
    _style_heading(h)

    clauses = risk_data.get('clauses', [])
    risk_clauses = [c for c in clauses if c.get('risk_level') in ('HIGH', 'MEDIUM')]
    low_clauses = [c for c in clauses if c.get('risk_level') == 'LOW']

    p = doc.add_paragraph()
    run = p.add_run(
        f'以下对合同全部 {len(clauses)} 条条款进行逐条分析。'
        f'其中高风险 {len([c for c in clauses if c.get("risk_level") == "HIGH"])} 条、'
        f'中风险 {len([c for c in clauses if c.get("risk_level") == "MEDIUM"])} 条、'
        f'低风险 {len(low_clauses)} 条。'
        f'每条风险项均标注了对应的法律依据和具体修改建议。'
    )
    run.font.size = Pt(10)
    doc.add_paragraph()

    if not risk_clauses and not low_clauses:
        p = doc.add_paragraph()
        p.add_run('未发现任何条款风险。')
        return

    # Section divider
    if risk_clauses:
        p = doc.add_paragraph()
        run = p.add_run('▬▬▬▬▬▬▬▬▬▬▬▬ 高风险与中风险条款 ▬▬▬▬▬▬▬▬▬▬▬▬')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(180, 0, 0)
        run.bold = True
        doc.add_paragraph()

        for clause in risk_clauses:
            _render_clause_detail(doc, clause)
            doc.add_paragraph()

    # Low risk clauses — summary only
    if low_clauses:
        p = doc.add_paragraph()
        run = p.add_run('▬▬▬▬▬▬▬▬▬▬▬▬ 低风险条款（摘要） ▬▬▬▬▬▬▬▬▬▬▬▬')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 128, 0)
        run.bold = True
        doc.add_paragraph()

        for clause in low_clauses:
            _render_low_risk_clause(doc, clause)
            doc.add_paragraph()


def _render_low_risk_clause(doc, clause):
    """Render a low-risk clause in compact summary format"""
    p = doc.add_paragraph()
    run = p.add_run(f"{clause.get('number', '')}  {clause.get('title', '条款')[:80]}")
    run.bold = True
    run.font.size = Pt(10)

    run = p.add_run(f"  [低风险 | 评分: {clause.get('risk_score', 0)}]")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 128, 0)

    # Brief content preview
    content = clause.get('content', '')
    if len(content) > 200:
        content = content[:200] + '...'
    p = doc.add_paragraph()
    run = p.add_run(f'{content}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(120, 120, 120)


def _render_clause_detail(doc, clause):
    level = clause.get('risk_level', 'LOW')
    level_colors = {
        'HIGH': RGBColor(200, 40, 40),
        'MEDIUM': RGBColor(200, 140, 0),
        'LOW': RGBColor(0, 128, 0),
    }
    level_bg = {
        'HIGH': 'FFD6D6',
        'MEDIUM': 'FFF3D6',
        'LOW': 'D6FFD6',
    }
    lc = level_colors.get(level, RGBColor(0, 0, 0))

    # ── Section separator ──
    p = doc.add_paragraph()
    run = p.add_run('─' * 50)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(200, 200, 200)

    # ── Title line with risk badge ──
    p = doc.add_paragraph()
    run = p.add_run(f"▎ {clause.get('number', '')}  {clause.get('title', '条款')[:100]}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = lc

    # Risk metadata line
    p = doc.add_paragraph()
    run = p.add_run(f'风险等级：{level}    风险评分：{clause.get("risk_score", 0)} 分    条款类别：{clause.get("category", "未分类")}')
    run.font.size = Pt(9)
    run.font.color.rgb = lc

    # ── Clause content (original text) ──
    content = clause.get('content', '')
    if len(content) > 1500:
        content = content[:1500] + '\n\n... [条款内容较长，完整内容请参见原始合同]'
    p = doc.add_paragraph()
    run = p.add_run('【条款原文】')
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(60, 60, 60)
    p = doc.add_paragraph()
    run = p.add_run(content)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(80, 80, 80)

    # Risk factors — show all, no artificial limit
    factors = clause.get('risk_factors', [])
    if factors:
        p = doc.add_paragraph()
        run = p.add_run('风险识别：')
        run.bold = True
        run.font.size = Pt(10)
        for f in factors:
            run = p.add_run(f'\n  • {f}')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(180, 0, 0)

    # Legal references — show all, no artificial limit
    refs = clause.get('legal_references', [])
    if refs:
        p = doc.add_paragraph()
        run = p.add_run('法律依据：')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 51, 102)
        for r in refs:
            run = p.add_run(f'\n  📖 {r}')
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0, 51, 102)

    # Suggestions — show all, no artificial limit
    suggestions = clause.get('suggestions', [])
    if suggestions:
        p = doc.add_paragraph()
        run = p.add_run('修改建议：')
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 100, 0)
        for s in suggestions:
            run = p.add_run(f'\n  ✓ {s}')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0, 100, 0)


# ══════════════════════════════════════════════════
# Civil Code Article Reference Table
# ══════════════════════════════════════════════════
def add_civil_code_reference_table(doc, clauses):
    h = doc.add_heading('七、《民法典》第463-647条条款覆盖清单', 1)
    _style_heading(h)

    p = doc.add_paragraph()
    run = p.add_run(
        '下表列出本次合同审查中实际涉及比对分析的《民法典》条款。'
        '具备法律依据定位的风险项均标注了对应的法条编号和条文内容。'
    )
    run.font.size = Pt(10)
    doc.add_paragraph()

    # Collect unique article references
    article_data = {}
    for c in clauses:
        ref = c.get('article_ref', '')
        if ref:
            if ref not in article_data:
                article_data[ref] = {'count': 0, 'max_risk': '', 'clauses': []}
            article_data[ref]['count'] += 1
            article_data[ref]['clauses'].append(c.get('title', '')[:60])
            risk = c.get('risk_level', '')
            if risk == 'HIGH':
                article_data[ref]['max_risk'] = 'HIGH'
            elif risk == 'MEDIUM' and article_data[ref]['max_risk'] != 'HIGH':
                article_data[ref]['max_risk'] = 'MEDIUM'
            elif article_data[ref]['max_risk'] == '':
                article_data[ref]['max_risk'] = 'LOW'

    if article_data:
        table = doc.add_table(rows=len(article_data) + 1, cols=4)
        table.style = 'Table Grid'

        headers = ['法条', '涉及条款数', '最高风险', '相关条款']
        for i, hdr in enumerate(headers):
            c = table.rows[0].cells[i]
            c.text = hdr
            for p in c.paragraphs:
                if p.runs:
                    p.runs[0].bold = True

        sorted_articles = sorted(article_data.items(), key=lambda x: x[1]['count'], reverse=True)
        for i, (art, data) in enumerate(sorted_articles[:40], 1):
            if i >= len(table.rows):
                break
            table.rows[i].cells[0].text = art
            table.rows[i].cells[1].text = str(data['count'])
            table.rows[i].cells[2].text = data['max_risk']
            table.rows[i].cells[3].text = '; '.join(data['clauses'][:3])


# ══════════════════════════════════════════════════
# Priority Recommendations
# ══════════════════════════════════════════════════
def add_priority_recommendations(doc, risk_data):
    h = doc.add_heading('八、优先修订建议', 1)
    _style_heading(h)

    summary = risk_data.get('summary', {})
    hc = summary.get('high_risk_count', 0)
    mc = summary.get('medium_risk_count', 0)

    # Priority 1: High risk MUST fix
    if hc > 0:
        p = doc.add_paragraph()
        run = p.add_run('【优先一】必须修订的高风险条款')
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(180, 0, 0)

        high_risk = risk_data.get('high_risk_clauses', [])
        for i, clause in enumerate(high_risk, 1):
            p = doc.add_paragraph()
            run = p.add_run(f'  {i}. {clause.get("title", "条款")[:80]}')
            run.bold = True
            run.font.size = Pt(10)

            suggestions = clause.get('suggestions', [])
            for s in suggestions[:2]:
                p2 = doc.add_paragraph()
                run = p2.add_run(f'    → {s}')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 100, 0)

        doc.add_paragraph()

    # Priority 2: Medium risk SHOULD negotiate
    if mc > 0:
        p = doc.add_paragraph()
        run = p.add_run('【优先二】建议协商的中风险条款')
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(200, 150, 0)

        medium_risk = risk_data.get('medium_risk_clauses', [])
        for i, clause in enumerate(medium_risk[:8], 1):
            p = doc.add_paragraph()
            run = p.add_run(f'  {i}. {clause.get("title", "条款")[:80]}')
            run.font.size = Pt(10)

            suggestions = clause.get('suggestions', [])
            for s in suggestions[:1]:
                p2 = doc.add_paragraph()
                run = p2.add_run(f'    → {s}')
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0, 100, 0)

        doc.add_paragraph()

    # General recommendations
    p = doc.add_paragraph()
    run = p.add_run('【优先三】通用建议')
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0, 51, 102)

    general = [
        '签署合同前请专业律师进行全面审核，确保所有条款符合现行法律法规。',
        '确保双方对所有条款理解一致，格式条款应主动向对方说明其内容及法律后果。',
        '保留合同谈判过程中的全部沟通记录（邮件、会议纪要、修订版本等）。',
        '明确约定违约责任的具体计算方式和争议解决方式，避免后续争议成本过高。',
        '保密条款应明确信息范围、保密期限（不超过2-5年）和违约救济措施。',
        '竞业限制期限不超过2年，必须约定经济补偿并明确支付方式。',
        '核心条款（标的、价款、履行方式、违约责任）不应依赖法定补充规则。',
        '定期审查合同的执行情况和外部环境变化，及时通过补充协议完善。',
    ]
    for g in general:
        p = doc.add_paragraph()
        run = p.add_run(f'  • {g}')
        run.font.size = Pt(10)


# ══════════════════════════════════════════════════
# Legal Reference Index
# ══════════════════════════════════════════════════
def add_legal_reference_index(doc, all_refs):
    h = doc.add_heading('九、法律依据索引', 1)
    _style_heading(h)

    # Group by law
    grouped = {}
    for ref in all_refs:
        if '民法典第' in ref or '民法典》第' in ref:
            key = '《中华人民共和国民法典》合同编'
        elif '商业秘密' in ref or '反不正当竞争' in ref:
            key = '《商业秘密保护规定》及相关法规'
        elif '民事诉讼法' in ref:
            key = '《中华人民共和国民事诉讼法》'
        else:
            key = '其他法律依据'
        if key not in grouped:
            grouped[key] = []
        grouped[key].append(ref)

    for law, refs in grouped.items():
        p = doc.add_paragraph()
        run = p.add_run(law)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102)

        for ref in sorted(set(refs))[:20]:
            p = doc.add_paragraph()
            run = p.add_run(f'  • {ref}')
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(80, 80, 80)

    # Additional reference laws
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('其他相关法律法规：')
    run.bold = True
    run.font.size = Pt(11)

    additional = [
        '《中华人民共和国民法典》（合同编第463-647条）',
        '《中华人民共和国反不正当竞争法》（2019年修订）',
        '最高人民法院《关于审理侵犯商业秘密民事案件适用法律若干问题的规定》（法释〔2020〕7号）',
        '最高人民法院《关于适用〈中华人民共和国民法典〉合同编通则若干问题的解释》',
        '《中华人民共和国电子签名法》',
        '《中华人民共和国民事诉讼法》',
        '《中华人民共和国仲裁法》',
    ]
    for ref in additional:
        p = doc.add_paragraph()
        run = p.add_run(f'  • {ref}')
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)


# ══════════════════════════════════════════════════
# Disclaimer
# ══════════════════════════════════════════════════
def add_disclaimer_final(doc):
    h = doc.add_heading('十、免责声明', 1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(12)

    text = (
        '本报告为自动化风险分析工具生成的参考性文件，不构成法律意见。'
        '报告的审查结果基于对合同文本的模式匹配和法律法规的普遍性解读，'
        '不能替代专业律师的个案判断。每一条风险的法律后果应在具体事实情境下'
        '由具有执业资格的中国律师进行评估。'
    )
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    text2 = (
        '风险评分和等级划分为基于预设规则的计算结果，作为风险程度的参考性评估。'
        '合同的效力、解释和履行最终应遵循相关法律法规的规定，'
        '由司法机关或仲裁机构根据完整的案件证据依法认定。'
        '建议在签署任何具有重大利益的合同前，咨询持有中华人民共和国执业资格的专业律师。'
    )
    p = doc.add_paragraph()
    run = p.add_run(text2)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)


# ══════════════════════════════════════════════════
# Footer
# ══════════════════════════════════════════════════
def add_footer(doc):
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('合同风险审查报告  |  基于《民法典》第463-647条及《商业秘密保护规定》  |  仅供参考不构成法律意见')
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(128, 128, 128)


def _style_heading(heading):
    """Apply consistent styling to a heading"""
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)


# ══════════════════════════════════════════════════
# Entry point
# ══════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(description='Generate comprehensive Word report')
    parser.add_argument('risk_data', help='Path to risk analysis JSON file')
    parser.add_argument('-o', '--output', help='Output Word document path')
    args = parser.parse_args()
    return generate_report(args.risk_data, args.output)


if __name__ == "__main__":
    sys.exit(main())

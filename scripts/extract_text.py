#!/usr/bin/env python3
"""Extract text from contract files (PDF/DOCX/TXT) for risk analysis."""

import sys
import os
import argparse


def extract_from_docx(filepath):
    from docx import Document
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Extract tables too
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return '\n'.join(parts)


def extract_from_txt(filepath):
    with open(filepath, 'r', encoding='utf-8') as f:
        return f.read()


def extract_from_pdf(filepath):
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            parts = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    parts.append(text)
            return '\n'.join(parts)
    except ImportError:
        pass
    try:
        import fitz
        doc = fitz.open(filepath)
        parts = []
        for page in doc:
            text = page.get_text()
            if text:
                parts.append(text)
        return '\n'.join(parts)
    except ImportError:
        raise RuntimeError("No PDF library available. Install pdfplumber or PyMuPDF.")


def main():
    parser = argparse.ArgumentParser(description='Extract text from contract files')
    parser.add_argument('input', help='Input file path (.docx/.pdf/.txt)')
    parser.add_argument('-o', '--output', help='Output text file path')
    args = parser.parse_args()

    ext = os.path.splitext(args.input)[1].lower()

    if ext == '.docx':
        text = extract_from_docx(args.input)
    elif ext == '.pdf':
        text = extract_from_pdf(args.input)
    elif ext == '.txt':
        text = extract_from_txt(args.input)
    else:
        print(f"ERROR: Unsupported format: {ext}")
        return 1

    text = text.strip()

    if args.output:
        with open(args.output, 'w', encoding='utf-8') as f:
            f.write(text)
        print(f"[OK] Extracted {len(text)} chars → {args.output}")
    else:
        print(text)

    return 0


if __name__ == "__main__":
    sys.exit(main())
